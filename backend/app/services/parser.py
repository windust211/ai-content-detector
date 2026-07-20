"""文档解析服务 — 从文件内容中提取纯文本。"""

import io
import os


def extract_text(content: bytes, ext: str) -> str:
    """根据文件内容（bytes）提取文本。适用于小文件。"""
    ext = ext.lower()

    if ext == ".txt":
        return _extract_from_txt(content)
    elif ext == ".pdf":
        return _extract_from_pdf(content)
    elif ext == ".docx":
        return _extract_from_docx(content)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def extract_text_from_file(file_path: str, ext: str) -> str:
    """根据文件路径提取文本。适用于大文件，避免全部加载到内存。"""
    ext = ext.lower()

    if ext == ".txt":
        return _extract_from_txt_file(file_path)
    elif ext == ".pdf":
        return _extract_from_pdf_file(file_path)
    elif ext == ".docx":
        return _extract_from_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def _extract_from_txt(content: bytes) -> str:
    """从纯文本 bytes 中提取文本。"""
    return content.decode("utf-8", errors="replace")


def _extract_from_txt_file(file_path: str) -> str:
    """从纯文本文件中分块读取提取文本。"""
    chunks = []
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        while chunk := f.read(8192):
            chunks.append(chunk)
    return "".join(chunks)


def _extract_from_pdf(content: bytes) -> str:
    """从 PDF bytes 中提取文本（依赖 PyMuPDF）。"""
    import fitz  # PyMuPDF
    doc = fitz.open(stream=content, filetype="pdf")
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return text


def _extract_from_pdf_file(file_path: str) -> str:
    """从 PDF 文件路径提取文本（依赖 PyMuPDF，直接读文件更高效）。"""
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return text


def _extract_from_docx(content: bytes) -> str:
    """从 DOCX bytes 中提取文本（依赖 python-docx）。"""
    from docx import Document
    doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_from_docx_file(file_path: str) -> str:
    """从 DOCX 文件路径提取文本（依赖 python-docx）。"""
    from docx import Document
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


def get_file_size(file_path: str) -> int:
    """获取文件大小（字节）。"""
    return os.path.getsize(file_path)